"""
PDF/DOCX 解析器封装模块

职责：将 PDF/DOCX 文档解析为统一的 Node JSON 序列。
使用 PyMuPDF (fitz) + pdfplumber 作为解析引擎
Node JSON 协议与 Java 侧 NodeDTO 完全对应

Node 类型：
- heading: 标题（基于字体大小/样式启发式识别）
- text: 普通文本段落
- image: 图片（提取并保存到 storePath）
- code: 代码块（基于字体/缩进启发式识别）
- table: 表格（pdfplumber 表格识别）
- formula: 公式（基于 LaTeX/数学符号启发式识别）
"""

import os
import hashlib
import logging
import re
import tempfile
from pathlib import Path
from typing import Optional, List, Dict, Any, Tuple

import fitz  # PyMuPDF
import pdfplumber
from PIL import Image

logger = logging.getLogger("docling_analysis_service.parser")


def _build_title_path(title_stack: List[tuple]) -> Optional[str]:
    """根据标题栈构建 titlePath（"一级 > 二级"），栈空返回 None。"""
    if not title_stack:
        return None
    return " > ".join(title for _, title in title_stack)


class DocumentParser:
    """文档解析器，基于 PyMuPDF + pdfplumber"""

    def __init__(
        self,
        image_store_dir: str,
        image_url_prefix: str = "/chunk_images",
    ):
        """
        :param image_store_dir: 图片存储根目录（Java 的 storePath/chunk_images）
        :param image_url_prefix: 图片 URL 前缀
        """
        self.image_store_dir = Path(image_store_dir)
        self.image_url_prefix = image_url_prefix.rstrip("/")
        logger.info(
            "DocumentParser 初始化完成，图片存储目录: %s",
            self.image_store_dir,
        )

    def parse(self, file_path: str, document_id: int, user_id: int) -> dict:
        """
        解析文档，返回统一 Node JSON

        :param file_path: 文档本地路径
        :param document_id: 文档 ID（用于图片目录隔离）
        :param user_id: 用户 ID（用于图片目录隔离）
        :return: {"documentType": "pdf", "nodes": [...]}
        """
        file_path = Path(file_path)
        document_type = self._detect_document_type(file_path)

        logger.info("开始解析文档: %s (type=%s)", file_path, document_type)

        # 非结构化文档（md/html/code/linebased）交由 routers 路由分发，
        # 由对应 Parser 完成结构识别 + 超长拆分 + titlePath/parentId 标注。
        if document_type not in ("pdf", "docx", "xlsx", "unknown"):
            from parsers.router import parse as router_parse
            return router_parse(file_path, document_id, user_id)

        # 图片输出目录: storePath/chunk_images/{userId}/{docId}/
        image_output_dir = self.image_store_dir / str(user_id) / str(document_id)
        image_output_dir.mkdir(parents=True, exist_ok=True)

        if document_type == "pdf":
            nodes = self._parse_pdf(file_path, image_output_dir, user_id, document_id)
        elif document_type == "docx":
            nodes = self._parse_docx(file_path, image_output_dir, user_id, document_id)
        else:
            logger.warning("不支持的文档类型: %s，返回空 Node 列表", document_type)
            nodes = []

        logger.info("文档解析完成，生成 %d 个 Node", len(nodes))
        return {"documentType": document_type, "nodes": nodes}

    # ------------------------------ PDF 解析 ------------------------------

    def _parse_pdf(
        self,
        file_path: Path,
        image_output_dir: Path,
        user_id: int,
        document_id: int,
    ) -> List[dict]:
        """使用 PyMuPDF + pdfplumber 解析 PDF"""
        nodes = []
        node_index = 0

        # 标题栈：跨页维护 PDF 的标题层级，推导 titlePath（与 docx parser 一致）
        title_stack: List[tuple] = []

        # PyMuPDF 打开文档
        doc = fitz.open(str(file_path))
        total_pages = doc.page_count
        logger.info("PDF 共 %d 页", total_pages)

        # pdfplumber 打开文档（用于表格识别）
        with pdfplumber.open(str(file_path)) as pdf_plumber:
            for page_num in range(total_pages):
                page = doc[page_num]
                plumber_page = pdf_plumber.pages[page_num] if page_num < len(pdf_plumber.pages) else None

                # 1. 提取图片
                image_nodes = self._extract_images_from_page(
                    page, page_num + 1, image_output_dir, user_id, document_id, node_index,
                    title_stack,
                )
                for node in image_nodes:
                    nodes.append(node)
                    node_index += 1

                # 2. 提取表格（pdfplumber）
                if plumber_page:
                    table_nodes = self._extract_tables_from_page(
                        plumber_page, page_num + 1, node_index, title_stack,
                    )
                    for node in table_nodes:
                        nodes.append(node)
                        node_index += 1

                # 3. 提取文本块（PyMuPDF，按阅读顺序）
                text_nodes = self._extract_text_blocks_from_page(
                    page, page_num + 1, node_index, title_stack,
                )
                for node in text_nodes:
                    nodes.append(node)
                    node_index += 1

        doc.close()
        return nodes

    def _extract_images_from_page(
        self,
        page: fitz.Page,
        page_num: int,
        image_output_dir: Path,
        user_id: int,
        document_id: int,
        start_index: int,
        title_stack: List[tuple],
    ) -> List[dict]:
        """从 PDF 页面提取图片"""
        nodes = []
        image_list = page.get_images(full=True)

        for img_index, img_info in enumerate(image_list):
            try:
                xref = img_info[0]
                base_image = page.parent.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]

                # 保存图片
                img_filename = f"img_{start_index + img_index + 1:03d}.{image_ext}"
                img_path = image_output_dir / img_filename
                with open(img_path, "wb") as f:
                    f.write(image_bytes)

                # 获取图片尺寸
                width, height = None, None
                try:
                    with Image.open(img_path) as pil_img:
                        width, height = pil_img.size
                except Exception:
                    pass

                # 计算图片哈希
                img_hash = self._compute_hash(image_bytes)

                # 构建相对 URL
                relative_url = f"{self.image_url_prefix}/{user_id}/{document_id}/{img_filename}"

                # 获取图片在页面上的位置（bbox）
                bbox = None
                try:
                    # 通过图片引用获取位置信息
                    for block in page.get_text("dict")["blocks"]:
                        if block.get("type") == 1:  # image block
                            bbox = [
                                block["bbox"][0],
                                block["bbox"][1],
                                block["bbox"][2] - block["bbox"][0],
                                block["bbox"][3] - block["bbox"][1],
                            ]
                            break
                except Exception:
                    pass

                nodes.append({
                    "id": f"n{start_index + img_index + 1}",
                    "type": "image",
                    "imagePath": relative_url,
                    "titlePath": _build_title_path(title_stack),
                    "page": page_num,
                    "bbox": bbox,
                    "width": width,
                    "height": height,
                    "hash": img_hash,
                })
                logger.debug("提取图片: %s (page=%d)", img_filename, page_num)

            except Exception as e:
                logger.warning("图片提取失败 (xref=%d): %s", img_info[0], e)

        return nodes

    def _extract_tables_from_page(
        self,
        page: pdfplumber.page.Page,
        page_num: int,
        start_index: int,
        title_stack: List[tuple],
    ) -> List[dict]:
        """使用 pdfplumber 提取表格"""
        nodes = []
        try:
            tables = page.extract_tables()
            for table_index, table_data in enumerate(tables):
                if not table_data:
                    continue

                # 转换为 HTML
                html = self._table_to_html(table_data)
                row_count = len(table_data)
                col_count = max(len(row) for row in table_data) if table_data else 0

                # 表格 bbox（pdfplumber 提供）
                bbox = None
                try:
                    table_bbox = page.find_tables()[table_index].bbox if table_index < len(page.find_tables()) else None
                    if table_bbox:
                        bbox = [table_bbox[0], table_bbox[1], table_bbox[2] - table_bbox[0], table_bbox[3] - table_bbox[1]]
                except Exception:
                    pass

                nodes.append({
                    "id": f"n{start_index + table_index + 1}",
                    "type": "table",
                    "html": html,
                    "rowCount": row_count,
                    "colCount": col_count,
                    "titlePath": _build_title_path(title_stack),
                    "page": page_num,
                    "bbox": bbox,
                })
                logger.debug("提取表格: %d行%d列 (page=%d)", row_count, col_count, page_num)

        except Exception as e:
            logger.warning("表格提取失败 (page=%d): %s", page_num, e)

        return nodes

    def _extract_text_blocks_from_page(
        self,
        page: fitz.Page,
        page_num: int,
        start_index: int,
        title_stack: List[tuple],
    ) -> List[dict]:
        """提取文本块，基于字体大小识别标题和正文"""
        nodes = []

        # 获取文本块（保留顺序）
        blocks = page.get_text("dict", sort=True)["blocks"]

        for block_index, block in enumerate(blocks):
            if block.get("type") != 0:  # 跳过图片块
                continue

            # 提取文本
            text = ""
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text += span.get("text", "")

            if not text.strip():
                continue

            # 分析字体特征
            is_heading = False
            heading_level = 1
            is_code = False

            try:
                spans = block.get("lines", [])[0].get("spans", []) if block.get("lines") else []
                if spans:
                    # 获取字体大小
                    font_size = spans[0].get("size", 12)
                    font_flags = spans[0].get("flags", 0)

                    # 标题启发式：大字体或粗体
                    if font_size > 14 or (font_flags & 16):  # 16 = bold
                        is_heading = True
                        heading_level = 1 if font_size > 18 else 2

                    # 代码启发式：等宽字体或特定字体名
                    font_name = spans[0].get("font", "").lower()
                    if any(mono in font_name for mono in ["mono", "consol", "courier", "code"]):
                        is_code = True

            except Exception:
                pass

            # 代码块额外启发式：连续空格/缩进
            if not is_code:
                if re.search(r"^\s{4,}", text) or text.count("    ") >= 3:
                    is_code = True

            # bbox
            bbox = [
                block["bbox"][0],
                block["bbox"][1],
                block["bbox"][2] - block["bbox"][0],
                block["bbox"][3] - block["bbox"][1],
            ]

            node_id = f"n{start_index + block_index + 1}"

            if is_heading:
                nodes.append({
                    "id": node_id,
                    "type": "heading",
                    "text": text.strip(),
                    "level": heading_level,
                    "page": page_num,
                    "bbox": bbox,
                })
            elif is_code:
                nodes.append({
                    "id": node_id,
                    "type": "code",
                    "text": text.strip(),
                    "language": self._detect_language(text),
                    "page": page_num,
                    "bbox": bbox,
                })
            else:
                nodes.append({
                    "id": node_id,
                    "type": "text",
                    "text": text.strip(),
                    "page": page_num,
                    "bbox": bbox,
                })

        return nodes

    # ------------------------------ DOCX 解析 ------------------------------

    def _parse_docx(
        self,
        file_path: Path,
        image_output_dir: Path,
        user_id: int,
        document_id: int,
    ) -> List[dict]:
        """解析 DOCX 文档，按阅读顺序提取文本、图片、表格"""
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(str(file_path))
        nodes = []
        node_index = 0
        image_counter = 1

        # 标题栈：维护 docx 的标题层级，推导 titlePath（与 markdown/html parser 一致）
        # [(level, title), ...]，遇新标题时弹出 level >= 当前的，保留上级
        title_stack: List[tuple] = []

        # 建立段落/表格元素到对象的快速映射
        paragraph_map = {id(p._element): p for p in doc.paragraphs}
        table_map = {id(t._element): t for t in doc.tables}

        for element in doc.element.body:
            tag = element.tag
            if tag == qn("w:p"):
                paragraph = paragraph_map.get(id(element))
                if paragraph is None:
                    continue

                text = paragraph.text.strip()

                # 提取段落中的内嵌图片（保持阅读顺序）
                # python-docx 的 xpath 已内置常见命名空间前缀，直接使用 a:blip / r:embed
                blip_elements = element.xpath(".//a:blip")
                for blip in blip_elements:
                    embed = blip.get(qn("r:embed"))
                    if not embed:
                        continue
                    image_part = doc.part.related_parts.get(embed)
                    if image_part is None:
                        continue
                    image_bytes = image_part.blob
                    ext = self._guess_image_ext(image_part.content_type)
                    img_filename = f"img_{image_counter:03d}.{ext}"
                    image_counter += 1
                    img_path = image_output_dir / img_filename
                    with open(img_path, "wb") as f:
                        f.write(image_bytes)

                    relative_url = f"{self.image_url_prefix}/{user_id}/{document_id}/{img_filename}"
                    nodes.append({
                        "id": f"n{node_index + 1}",
                        "type": "image",
                        "imagePath": relative_url,
                        "titlePath": _build_title_path(title_stack),
                        "page": 1,
                        "bbox": None,
                    })
                    node_index += 1

                if not text:
                    continue

                # 判断标题样式
                style_name = paragraph.style.name if paragraph.style else ""
                is_heading = style_name.startswith("Heading")
                level = 1
                if is_heading and len(style_name) > 8:
                    try:
                        level = int(style_name.split()[-1])
                    except ValueError:
                        level = 1

                # 当前 titlePath（push 前的栈状态对应当前块所属标题路径）
                current_title_path = _build_title_path(title_stack)

                if is_heading:
                    # 更新标题栈：弹出 level >= 当前的，压入当前标题
                    while title_stack and title_stack[-1][0] >= level:
                        title_stack.pop()
                    title_stack.append((level, text))

                    nodes.append({
                        "id": f"n{node_index + 1}",
                        "type": "heading",
                        "text": text,
                        "level": level,
                        "titlePath": _build_title_path(title_stack),
                        "page": 1,
                        "bbox": None,
                    })
                else:
                    # 简单代码启发式：以 4 个空格开头或包含大量缩进
                    is_code = text.startswith("    ") or text.count("    ") >= 2
                    language = self._detect_language(text) if is_code else None
                    nodes.append({
                        "id": f"n{node_index + 1}",
                        "type": "code" if is_code else "text",
                        "text": text,
                        "language": language,
                        "titlePath": current_title_path,
                        "page": 1,
                        "bbox": None,
                    })
                node_index += 1

            elif tag == qn("w:tbl"):
                table = table_map.get(id(element))
                if table is None:
                    continue
                html = self._docx_table_to_html(table)
                row_count = len(table.rows)
                col_count = max(len(row.cells) for row in table.rows) if table.rows else 0
                nodes.append({
                    "id": f"n{node_index + 1}",
                    "type": "table",
                    "html": html,
                    "rowCount": row_count,
                    "colCount": col_count,
                    "titlePath": _build_title_path(title_stack),
                    "page": 1,
                    "bbox": None,
                })
                node_index += 1

        logger.info("DOCX 解析完成，生成 %d 个 Node", len(nodes))
        return nodes

    def _docx_table_to_html(self, table) -> str:
        """将 DOCX 表格转换为 HTML"""
        html_parts = ['<table border="1">']
        for i, row in enumerate(table.rows):
            html_parts.append("<tr>")
            for cell in row.cells:
                tag = "th" if i == 0 else "td"
                html_parts.append(f"<{tag}>{cell.text}</{tag}>")
            html_parts.append("</tr>")
        html_parts.append("</table>")
        return "".join(html_parts)

    def _guess_image_ext(self, content_type: str) -> str:
        """根据 MIME 类型猜测图片扩展名"""
        mapping = {
            "image/png": "png",
            "image/jpeg": "jpg",
            "image/gif": "gif",
            "image/bmp": "bmp",
            "image/webp": "webp",
        }
        return mapping.get(content_type, "png")

    # ------------------------------ 辅助方法 ------------------------------

    def _detect_document_type(self, file_path: Path) -> str:
        """根据扩展名检测文档类型"""
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return "pdf"
        if suffix in (".docx", ".doc"):
            return "docx"
        if suffix == ".xlsx":
            return "xlsx"
        return "unknown"

    def _table_to_html(self, table_data: List[List[str]]) -> str:
        """将表格数据转换为 HTML"""
        if not table_data:
            return ""

        html_parts = ['<table border="1">']
        for i, row in enumerate(table_data):
            tag = "th" if i == 0 else "td"
            cells = [f"<{tag}>{cell if cell else ''}</{tag}>" for cell in row]
            html_parts.append("<tr>" + "".join(cells) + "</tr>")
        html_parts.append("</table>")
        return "".join(html_parts)

    def _detect_language(self, code_text: str) -> Optional[str]:
        """简单启发式检测代码语言"""
        if not code_text:
            return None
        text = code_text[:200]
        if "public class" in text or "System.out" in text:
            return "java"
        if "def " in text or "import " in text:
            return "python"
        if "function " in text or "var " in text or "const " in text:
            return "javascript"
        if "SELECT " in text.upper() or "FROM " in text.upper():
            return "sql"
        return None

    def _compute_hash(self, data: bytes) -> str:
        """计算数据的 MD5 哈希"""
        return hashlib.md5(data).hexdigest()