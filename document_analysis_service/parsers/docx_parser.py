"""
DOCX 解析器

职责：将 DOCX 文档解析为统一的 Node JSON 序列。
引擎：python-docx，按 body 元素阅读顺序遍历 w:p（段落）与 w:tbl（表格）。
与 parsers 包内其它解析器（markdown / html / code / linebased / pdf）平行。

Node JSON 协议（与 Java 侧 NodeDTO 对应）：
- id: str "n1"... 自管递增
- type: "heading" | "text" | "image" | "code" | "table"
- text / imagePath / html / language / level / page(固定1) / bbox(None) / titlePath / groupId(超长聚类内部拆出的小 Node 共享同一 groupId；普通块 None)

核心策略：
1. 遍历 doc.element.body，按 w:p / w:tbl 顺序处理，保持原文阅读顺序
2. 段落标题识别：
   - 优先 style.name（"Heading N" / "heading N" / "标题 N" / "Title"）
   - outline level（w:outlineLvl N）兜底，覆盖"正文样式但大纲级别为标题"的情况
   - 维护标题栈 [(level, title)] 推导 titlePath（与 markdown / html parser 一致）
3. 段落内嵌图片：a:blip / r:embed 提取并落盘，保持阅读顺序
4. 表格：行 → HTML（首行 th），作为原子块
5. 代码启发式：等宽字体名 / 4 空格缩进 / 围栏符号
6. 正文段落聚类缓冲：
   - 相邻正文 w:p 不再各自独立成 Node，而是累积到 cluster_buffer（文本用 \n 连接）
   - flush 触发：空 w:p / space_after 显著大 / 标题 / 代码 / 表格 / 列表项 / body 遍历结束
   - 单换行、段内 w:br/w:cr 一律合并，绝不作为分隔符
   - w:br / w:cr 用 XPath 显式取，渲染为段内 \n（不依赖 paragraph.text 渲染行为）
"""

import logging
import re
from pathlib import Path
from typing import List, Optional, Tuple

from ._common import (
    build_title_path,
    compute_hash,
    detect_code_language,
    guess_image_ext,
    make_node_seq,
    should_flush_under_elastic,
    table_to_html,
)

logger = logging.getLogger("docling_analysis_service.parsers.docx_parser")

# 等宽字体名特征（与 pdf_parser 一致）
_MONO_FONT_HINTS = ("mono", "consol", "courier", "code", "menlo", "consolas", "jetbrains")

# 标题样式名 → level 的正则（兼容 Word 内置英文名、中文本地化名、小写）
# - "Heading 1" / "heading 1" / "Heading1"
# - "标题 1" / "标题1"
# - "Title" 视为 level 1
_HEADING_STYLE_PATTERN = re.compile(
    r"^(?:heading|标题)\s*([1-9]\d?)$", re.IGNORECASE
)

# 代码围栏启发式：以 ``` 开头
_CODE_FENCE_PATTERN = re.compile(r"^\s*```")

# 4 空格缩进代码启发式
_INDENT_CODE_PATTERN = re.compile(r"^ {4,}\S")

# space_after 硬分隔阈值（pt）：段落直接 space_after >= 此值视为硬分隔，切断聚类。
# 12pt = Word 常见"段后间距"档位（默认 Normal 通常 8-10pt），覆盖作者显式设大间距的场景。
_HARD_SEPARATOR_SPACE_AFTER_PT = 12

# 超长聚类拆分阈值（与 markdown_parser / html_parser 的 CHUNK_THRESHOLD 一致）
CHUNK_THRESHOLD = 1000

# 句子分隔符（中英文，与 markdown_parser SENTENCE_DELIMITER 一致）
_SENTENCE_DELIMITER = re.compile(r"[。！？.!?；;]")


class DocxParser:
    """DOCX 解析器，基于 python-docx。

    与其它 parsers 平行：parse(file_path, document_id, user_id) -> List[Node dict]。
    图片落盘到 self.image_store_dir/{user_id}/{document_id}/。
    """

    def __init__(
        self,
        image_store_dir: str,
        image_url_prefix: str = "/chunk_images",
        chunk_threshold: int = CHUNK_THRESHOLD,
    ):
        self.image_store_dir = Path(image_store_dir)
        self.image_url_prefix = image_url_prefix.rstrip("/")
        self.threshold = chunk_threshold
        logger.info("DocxParser 初始化完成，图片存储目录: %s, 阈值: %d", self.image_store_dir, self.threshold)

    # ============================== 对外入口 ==============================

    def parse(
        self, file_path: str, document_id: int, user_id: int
    ) -> List[dict]:
        """
        解析 DOCX，返回 Node JSON 列表。

        :param file_path: DOCX 本地路径
        :param document_id: 文档 ID（图片目录隔离）
        :param user_id: 用户 ID（图片目录隔离）
        :return: Node dict 列表
        """
        from docx import Document
        from docx.oxml.ns import qn

        file_path = Path(file_path)
        logger.info("DocxParser 开始解析: %s", file_path)

        image_output_dir = self.image_store_dir / str(user_id) / str(document_id)
        image_output_dir.mkdir(parents=True, exist_ok=True)

        doc = Document(str(file_path))
        nodes: List[dict] = []
        node_seq = make_node_seq()

        # 标题栈：[(level, title)]，跨段维护层级
        title_stack: List[Tuple[int, str]] = []

        # 正文聚类缓冲状态（2026-07-06 决策）：
        # 相邻正文 w:p 累积到 buffer（flush 时文本用 \n 连接），
        # title_path 取聚类内首个段落的（同属一个标题区间）。
        cluster = {"buffer": [], "title_path": None}  # type: dict

        def flush_cluster():
            """flush 聚类缓冲为 text Node（若有内容）。

            合并后超长（> self.threshold）时：内部拆分为多个小 Node，共享同一 groupId（同源整块），
            不再输出整块超长镜像父 Node；由 Java 侧 NodeBasedChunkBuilder 据 groupId 合成父子 chunk。
            """
            if not cluster["buffer"]:
                return
            merged_text = "\n".join(cluster["buffer"])
            title_path = cluster["title_path"]
            cluster["buffer"].clear()
            cluster["title_path"] = None

            if len(merged_text) <= self.threshold:
                # 正常产出单个 text Node（普通块，无 groupId）
                nodes.append({
                    "id": next(node_seq),
                    "type": "text",
                    "text": merged_text,
                    "titlePath": title_path,
                    "groupId": None,
                    "page": 1,
                    "bbox": None,
                })
            else:
                # 超长拆分：内部拆为多个小 Node，共享同一 groupId（仿 markdown_parser.py 超长段处理）
                group_id = next(node_seq)
                sub_chunks = self._split_oversized_cluster(merged_text, self.threshold)
                for sub_text in sub_chunks:
                    if sub_text.strip():
                        nodes.append({
                            "id": next(node_seq),
                            "type": "text",
                            "text": sub_text.strip(),
                            "titlePath": title_path,
                            "groupId": group_id,
                            "page": 1,
                            "bbox": None,
                        })

        # 元素 → 对象映射，按 body 顺序遍历时取回
        paragraph_map = {id(p._element): p for p in doc.paragraphs}
        table_map = {id(t._element): t for t in doc.tables}

        for element in doc.element.body:
            tag = element.tag
            if tag == qn("w:p"):
                paragraph = paragraph_map.get(id(element))
                if paragraph is None:
                    continue
                self._handle_paragraph(
                    paragraph, element, doc, image_output_dir,
                    user_id, document_id, node_seq, title_stack, nodes,
                    cluster, flush_cluster,
                )
            elif tag == qn("w:tbl"):
                table = table_map.get(id(element))
                if table is None:
                    continue
                # 表格是原子块，切断正文聚类
                flush_cluster()
                self._handle_table(
                    table, node_seq, title_stack, nodes,
                )

        # body 遍历结束，flush 末尾缓冲
        flush_cluster()

        logger.info("DocxParser 解析完成: %s, 共 %d 个 Node", file_path, len(nodes))
        return nodes

    # ============================== 段落处理 ==============================

    def _handle_paragraph(
        self,
        paragraph,
        element,
        doc,
        image_output_dir: Path,
        user_id: int,
        document_id: int,
        node_seq,
        title_stack: List[Tuple[int, str]],
        nodes: List[dict],
        cluster: dict,
        flush_cluster,
    ) -> None:
        """处理 w:p 段落：聚类缓冲模式。

        正文段落累积到 cluster 缓冲（文本用 \\n 连接）；
        遇空 w:p / space_after 显著大 / 标题 / 代码 / 表格 / 列表项 / 图片 时 flush 为独立 text Node。
        单换行、段内 w:br/w:cr 一律合并，绝不作为分隔符。
        """
        # 用 XPath 显式取 w:br / w:cr 作为段内 \n，不依赖 paragraph.text 渲染行为
        text = self._extract_paragraph_text_with_breaks(paragraph).strip()

        # 1. 内嵌图片（即使段落无文本也提取）——图片是原子块，先 flush 正文聚类
        if element.xpath(".//a:blip"):
            flush_cluster()
        self._extract_inline_images(
            element, doc, image_output_dir,
            user_id, document_id, node_seq, title_stack, nodes,
        )

        if not text:
            # 空 w:p = 硬分隔，flush 缓冲
            flush_cluster()
            return

        # 2. 判定标题：style.name 优先，outline level 兜底
        is_heading, level = self._detect_heading(paragraph)

        if is_heading:
            # 标题是原子块，先 flush 正文聚类
            flush_cluster()
            # 更新标题栈：弹出 level >= 当前的，压入当前
            while title_stack and title_stack[-1][0] >= level:
                title_stack.pop()
            title_stack.append((level, text))

            nodes.append({
                "id": next(node_seq),
                "type": "heading",
                "text": text,
                "level": level,
                "titlePath": build_title_path(title_stack),
                "page": 1,
                "bbox": None,
            })
            return

        # 3. 判定列表项（w:numPr 存在）——原子块，切断正文聚类
        if self._is_list_item(paragraph):
            flush_cluster()
            nodes.append({
                "id": next(node_seq),
                "type": "text",
                "text": text,
                "titlePath": build_title_path(title_stack),
                "page": 1,
                "bbox": None,
            })
            return

        # 4. 判定代码——原子块，切断正文聚类
        is_code, language = self._detect_code(paragraph, text)
        if is_code:
            flush_cluster()
            nodes.append({
                "id": next(node_seq),
                "type": "code",
                "text": text,
                "language": language,
                "titlePath": build_title_path(title_stack),
                "page": 1,
                "bbox": None,
            })
            return

        # 5. 正文段落——加入聚类缓冲
        if not cluster["buffer"]:
            # 首个段落，捕获 titlePath（同属一个标题区间）
            cluster["title_path"] = build_title_path(title_stack)
        cluster["buffer"].append(text)

        # 6. space_after 显著大——当前段加入缓冲后 flush（硬分隔）
        if self._has_significant_space_after(paragraph):
            flush_cluster()

    # ============================== 内嵌图片提取 ==============================

    def _extract_inline_images(
        self,
        element,
        doc,
        image_output_dir: Path,
        user_id: int,
        document_id: int,
        node_seq,
        title_stack: List[Tuple[int, str]],
        nodes: List[dict],
    ) -> None:
        """提取段落中的内嵌图片（a:blip / r:embed），保持阅读顺序落盘。"""
        # python-docx 的 xpath 已内置常见命名空间前缀，直接使用 a:blip
        blip_elements = element.xpath(".//a:blip")
        for blip in blip_elements:
            try:
                from docx.oxml.ns import qn
                embed = blip.get(qn("r:embed"))
                if not embed:
                    continue
                image_part = doc.part.related_parts.get(embed)
                if image_part is None:
                    continue
                image_bytes = image_part.blob
                ext = guess_image_ext(image_part.content_type)

                node_id = next(node_seq)
                img_filename = f"img_{node_id[1:]}.{ext}"
                img_path = image_output_dir / img_filename
                with open(img_path, "wb") as f:
                    f.write(image_bytes)

                relative_url = (
                    f"{self.image_url_prefix}/{user_id}/{document_id}/{img_filename}"
                )
                img_hash = compute_hash(image_bytes)

                nodes.append({
                    "id": node_id,
                    "type": "image",
                    "imagePath": relative_url,
                    "titlePath": build_title_path(title_stack),
                    "page": 1,
                    "bbox": None,
                    "hash": img_hash,
                })
            except Exception as e:
                logger.warning("DOCX 内嵌图片提取失败: %s", e)

    # ============================== 标题判定 ==============================

    @staticmethod
    def _detect_heading(paragraph) -> Tuple[bool, int]:
        """
        判定段落是否为标题及 level。
        优先级：style.name（Heading N / 标题 N / Title）→ outline level（w:outlineLvl）。
        修复旧实现的问题：
        - 旧实现 style_name.startswith("Heading") 漏掉小写 "heading" 与中文 "标题"
        - 旧实现 len(style_name) > 8 的 gate 让 "Heading"（无数字）落到 level=1 默认，
          但 "Heading 10" 等也能 parseInt，逻辑混乱
        """
        # 1. style.name
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name:
            level = DocxParser._parse_heading_style(style_name)
            if level is not None:
                return True, level
            # Word 的 "Title" 样式视为 level 1
            if style_name.lower() == "title":
                return True, 1

        # 2. outline level 兜底（正文样式但大纲级别为标题）
        level = DocxParser._outline_level(paragraph)
        if level is not None and 0 <= level <= 8:
            # w:outlineLvl 是 0-based，0 = level 1
            return True, level + 1

        return False, 1

    @staticmethod
    def _parse_heading_style(style_name: str) -> Optional[int]:
        """从样式名解析标题 level，非标题样式返回 None。"""
        m = _HEADING_STYLE_PATTERN.match(style_name.strip())
        if m:
            lvl = int(m.group(1))
            # level 限制在 1-9
            return max(1, min(9, lvl))
        return None

    @staticmethod
    def _outline_level(paragraph) -> Optional[int]:
        """从段落 XML 提取 w:outlineLvl 值，无则 None。"""
        try:
            from docx.oxml.ns import qn
            pPr = paragraph._element.find(qn("w:pPr"))
            if pPr is None:
                return None
            outline = pPr.find(qn("w:outlineLvl"))
            if outline is None:
                return None
            val = outline.get(qn("w:val"))
            if val is None:
                return None
            return int(val)
        except (ValueError, AttributeError):
            return None
        except Exception:
            return None

    # ============================== 代码判定 ==============================

    @staticmethod
    def _detect_code(paragraph, text: str) -> Tuple[bool, Optional[str]]:
        """
        代码启发式：
        1. 围栏 ``` 开头 → code（最可靠）
        2. 段落 run 字体名等宽 → code
        3. 4 空格缩进开头（与旧实现一致，但旧实现还把 text.count('    ')>=2 作为代码，
           这会误判含两个以上 4 连续空格的普通段落；改为只在"行首 4 空格"时判定）
        """
        # 1. 围栏
        if _CODE_FENCE_PATTERN.match(text):
            return True, detect_code_language(text)

        # 2. 等宽字体（任一 run 命中即视为代码块特征）
        for run in paragraph.runs:
            font_name = (run.font.name or "").lower()
            if font_name and any(m in font_name for m in _MONO_FONT_HINTS):
                return True, detect_code_language(text)

        # 3. 4 空格缩进行（按行首判断，避免正文多空格误判）
        if _INDENT_CODE_PATTERN.match(text):
            return True, detect_code_language(text)

        return False, None

    # ============================== 段落文本与结构信号 ==============================

    @staticmethod
    def _extract_paragraph_text_with_breaks(paragraph) -> str:
        """提取段落文本，w:br / w:cr 渲染为段内 \\n，w:tab 渲染为 \\t。

        不依赖 paragraph.text 的渲染行为（python-docx 对 w:br 行为不一致），
        统一以 XPath 显式取为准（2026-07-06 决策第二节第 2 条）。
        XPath union 结果按文档顺序返回，保证阅读顺序正确。
        """
        from docx.oxml.ns import qn
        parts: List[str] = []
        for node in paragraph._element.xpath(".//w:t | .//w:br | .//w:cr | .//w:tab"):
            tag = node.tag
            if tag == qn("w:t"):
                parts.append(node.text or "")
            elif tag == qn("w:br") or tag == qn("w:cr"):
                parts.append("\n")
            elif tag == qn("w:tab"):
                parts.append("\t")
        return "".join(parts)

    @staticmethod
    def _is_list_item(paragraph) -> bool:
        """检测段落是否为列表项（w:pPr/w:numPr 存在）。

        列表项保持原子性，不进正文聚类缓冲，且自身作为切断信号
        （2026-07-06 决策第三节第 3 条）。
        """
        try:
            from docx.oxml.ns import qn
            pPr = paragraph._element.find(qn("w:pPr"))
            if pPr is None:
                return False
            return pPr.find(qn("w:numPr")) is not None
        except Exception:
            return False

    @staticmethod
    def _has_significant_space_after(paragraph) -> bool:
        """检测段落的直接 space_after 是否显著大（视为硬分隔）。

        只检查段落直接格式（paragraph_format.space_after），不回溯样式继承——
        覆盖「作者显式设大段后间距」的场景；样式继承的默认间距不触发
        （2026-07-06 决策第二节第 3 条）。
        """
        try:
            space_after = paragraph.paragraph_format.space_after
            if space_after is None:
                return False
            return space_after.pt >= _HARD_SEPARATOR_SPACE_AFTER_PT
        except Exception:
            return False

    # ============================== 超长聚类拆分 ==============================
    # 用于聚类合并后超长文本的内部二次切分（拆为多个小 Node，调用方写同一 groupId）。
    # docx 聚类文本用 \n 连接原始段落，与 markdown 的空行分隔不同，
    # 故直接按 \n 拆分原始段落，而非复用 markdown_parser 的 _split_by_weak_paragraphs。

    def _split_oversized_cluster(self, text: str, threshold: int) -> List[str]:
        """拆分超长聚类文本：按 \\n 拆为原始段落，累加到阈值。

        单个段落仍超长（> threshold）则按句子兜底拆分。
        """
        paragraphs = [p for p in text.split("\n") if p.strip()]
        if not paragraphs:
            return []

        # 所有段落都 <= threshold → 累加到阈值
        if all(len(p) <= threshold for p in paragraphs):
            return self._accumulate_with_threshold(paragraphs, threshold)

        # 有超长段落 → 按句子拆分兜底
        return self._split_by_sentence_with_threshold(text, threshold)

    @staticmethod
    def _accumulate_with_threshold(chunks: List[str], threshold: int) -> List[str]:
        """将拆分后的子块累加到阈值后输出，子块间以单换行连接。
        累加采用弹性区间（减少人为切断强关联语义）。"""
        results: List[str] = []
        if not chunks:
            return results

        buffer = ""
        for chunk in chunks:
            if not chunk or not chunk.strip():
                continue
            trimmed = chunk.strip()
            add_len = len(trimmed) + (1 if buffer else 0)
            if should_flush_under_elastic(len(buffer), add_len, threshold):
                results.append(buffer.strip())
                buffer = trimmed
            else:
                if buffer:
                    buffer += "\n"
                buffer += trimmed

        if buffer:
            results.append(buffer.strip())

        return results

    def _split_by_sentence_with_threshold(self, text: str, threshold: int) -> List[str]:
        """按句子分隔符拆分文本，累加句子直到最接近阈值，不切断单个句子。
        累加采用弹性区间（减少对强关联句的人为切断）。"""
        results: List[str] = []
        if not text or not text.strip():
            return results

        sentences = self._split_into_sentences(text)

        buffer = ""
        for sentence in sentences:
            if not sentence.strip():
                continue
            add_len = len(sentence)
            if should_flush_under_elastic(len(buffer), add_len, threshold):
                results.append(buffer.strip())
                buffer = sentence
            else:
                buffer += sentence

        if buffer:
            results.append(buffer.strip())

        return results

    @staticmethod
    def _split_into_sentences(text: str) -> List[str]:
        """按句子分隔符拆分文本，保留分隔符在句子末尾。

        （与 markdown_parser._split_into_sentences / Java splitIntoSentences 一致）
        """
        sentences: List[str] = []
        if not text:
            return sentences

        last_end = 0
        for m in _SENTENCE_DELIMITER.finditer(text):
            delimiter_end = m.end()
            sentence = text[last_end:delimiter_end]
            if sentence.strip():
                sentences.append(sentence)
            last_end = delimiter_end

        if last_end < len(text):
            remaining = text[last_end:]
            if remaining.strip():
                sentences.append(remaining)

        return sentences

    # ============================== 表格处理 ==============================

    def _handle_table(
        self,
        table,
        node_seq,
        title_stack: List[Tuple[int, str]],
        nodes: List[dict],
    ) -> None:
        """将 DOCX 表格转为 HTML 原子块。"""
        # 二维化：python-docx 的 cell.text 已合并后代文本
        rows_data: List[List[Optional[str]]] = []
        for row in table.rows:
            rows_data.append([cell.text for cell in row.cells])

        html = table_to_html(rows_data)
        row_count = len(table.rows)
        col_count = max(len(row.cells) for row in table.rows) if table.rows else 0

        nodes.append({
            "id": next(node_seq),
            "type": "table",
            "html": html,
            "rowCount": row_count,
            "colCount": col_count,
            "titlePath": build_title_path(title_stack),
            "page": 1,
            "bbox": None,
        })

    # ============================== 扩展名 / 内容特征判定 ==============================

    @staticmethod
    def supports_extension(file_path: str) -> bool:
        """扩展名判定"""
        suffix = Path(file_path).suffix.lower()
        return suffix in (".docx", ".doc")
